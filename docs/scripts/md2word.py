# -*- coding: utf-8 -*-
"""
Convertit les protocoles Markdown (docs/protocoles/*.md) en fichiers Word
(.docx), adaptes au copier-coller dans OneNote / Note SharePoint.

Contrairement au .txt, le .docx conserve une vraie mise en forme :
  - Titre #/##/###     -> styles Heading 1/2/3
  - Gras **x**         -> run en gras
  - Code inline `x`    -> run monospace (Consolas)
  - Blocs de code ```  -> paragraphe monospace (police Consolas, gris)
  - Tableaux |...|     -> vraie table Word (style Table Grid)
  - Citations >        -> paragraphe indente en italique
  - Listes - / 1.      -> List Bullet / List Number
  - Liens [txt](url)   -> txt (url)
  - Accents preserves

Usage :
  python md2word.py [REPERTOIRE_SORTIE]

  - Sans argument : ecrit dans ./Protocoles_Word (a cote du script)
  - Avec argument : ecrit dans le repertoire fourni

Dependances : python-docx  (pip install python-docx)
"""
import os
import re
import sys
import glob

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_DIR = os.path.normpath(os.path.join(SCRIPT_DIR, '..', 'protocoles'))
OUT_DIR = sys.argv[1] if len(sys.argv) > 1 else os.path.join(SCRIPT_DIR, 'Protocoles_Word')

os.makedirs(OUT_DIR, exist_ok=True)

INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')
LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def resolve_links(text):
    return LINK_RE.sub(r'\1 (\2)', text)


def add_runs(paragraph, text):
    """Ajoute le texte au paragraphe en gerant gras (**) et code inline (`)."""
    text = resolve_links(text)
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, rows):
    parsed = []
    for r in rows:
        cells = [resolve_links(c.strip().replace('**', '').replace('`', ''))
                 for c in r.strip().strip('|').split('|')]
        parsed.append(cells)
    data = [c for c in parsed if not all(re.fullmatch(r':?-{2,}:?', x or '-') for x in c)]
    if not data:
        return
    ncol = max(len(r) for r in data)
    for r in data:
        while len(r) < ncol:
            r.append('')
    table = doc.add_table(rows=len(data), cols=ncol)
    table.style = 'Table Grid'
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(val)
            if ri == 0:
                run.bold = True
            run.font.size = Pt(9)
    doc.add_paragraph()


def convert(md, doc):
    lines = md.split('\n')
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
            in_code = not in_code
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # tableaux
        if line.strip().startswith('|') and i + 1 < len(lines) and re.search(r'\|.*-{2,}.*\|', lines[i + 1]):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i])
                i += 1
            add_table(doc, tbl)
            continue

        # titres
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = min(len(m.group(1)), 3)
            txt = resolve_links(m.group(2).replace('**', '').replace('`', '')).strip()
            doc.add_heading(txt, level=level)
            i += 1
            continue

        # regles horizontales
        if re.fullmatch(r'\s*([-*_])\1{2,}\s*', line):
            i += 1
            continue

        # citations >
        m = re.match(r'^\s*>\s?(.*)$', line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(resolve_links(m.group(1).replace('**', '').replace('`', '')))
            r.italic = True
            i += 1
            continue

        # listes a puces
        m = re.match(r'^\s*[-*]\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, m.group(1))
            i += 1
            continue

        # listes numerotees
        m = re.match(r'^\s*\d+\.\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(1))
            i += 1
            continue

        # ligne vide
        if line.strip() == '':
            i += 1
            continue

        # paragraphe courant
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1


def main():
    count = 0
    for path in sorted(glob.glob(os.path.join(SRC_DIR, '*.md'))):
        name = os.path.splitext(os.path.basename(path))[0]
        with open(path, 'r', encoding='utf-8') as f:
            md = f.read()
        doc = Document()
        # police de base lisible
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
        convert(md, doc)
        out_path = os.path.join(OUT_DIR, name + '.docx')
        doc.save(out_path)
        count += 1
        print(f"OK  {name}.docx")
    print(f"\nTotal : {count} fichiers Word generes dans {OUT_DIR}")


if __name__ == '__main__':
    main()
